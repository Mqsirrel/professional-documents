"""
docx_craft.py - Professional Document Styling & Engineering Engine for python-docx

Provides production-ready typography, layout geometry, table engineering,
callouts, dynamic page numbers, cover archetypes, and headless rendering.
"""

from __future__ import annotations
import os
import shutil
import subprocess
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Union

import docx
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches, Pt, RGBColor

# ---------------------------------------------------------------------------
# 1. CURATED COLOR PALETTES (Hex & RGB)
# ---------------------------------------------------------------------------

PALETTES: Dict[str, Dict[str, str]] = {
    "slate_executive": {
        "primary": "0F172A",      # Deep Navy / Slate 900
        "secondary": "1E293B",    # Slate 800
        "accent": "2563EB",       # Royal Blue
        "accent_light": "EFF6FF", # Subtle Blue Tint
        "text": "0F172A",         # Near Black
        "text_muted": "64748B",   # Slate 500
        "surface": "F8FAFC",      # Slate 50
        "border": "CBD5E1",       # Slate 300
        "border_subtle": "E2E8F0",# Slate 200
        "highlight": "F59E0B",    # Amber
    },
    "nordic_frost": {
        "primary": "1E293B",      # Dark Charcoal
        "secondary": "334155",    # Medium Charcoal
        "accent": "0284C7",       # Sky Blue 600
        "accent_light": "F0F9FF", # Sky 50
        "text": "1E293B",         # Deep Slate
        "text_muted": "64748B",   # Slate
        "surface": "F8FAFC",      # Off White
        "border": "CBD5E1",       # Border
        "border_subtle": "E2E8F0",# Subtle
        "highlight": "0D9488",    # Teal
    },
    "academic_crimson": {
        "primary": "7F1D1D",      # Deep Burgundy / Crimson
        "secondary": "991B1B",    # Dark Crimson
        "accent": "B91C1C",       # Crimson Accent
        "accent_light": "FEF2F2", # Crimson Tint
        "text": "1C1917",         # Warm Black
        "text_muted": "78716C",   # Warm Gray
        "surface": "FAF8F5",      # Warm Ivory / Parchment
        "border": "D6D3D1",       # Stone 300
        "border_subtle": "E7E5E4",# Stone 200
        "highlight": "D97706",    # Warm Amber
    },
    "forest_emerald": {
        "primary": "14532D",      # Deep Forest Green
        "secondary": "166534",    # Forest Green
        "accent": "059669",       # Emerald
        "accent_light": "ECFDF5", # Mint Tint
        "text": "0F172A",         # Slate Black
        "text_muted": "64748B",   # Muted
        "surface": "F7F9F6",      # Soft Sage Tint
        "border": "CBD5E1",       # Border
        "border_subtle": "E2E8F0",# Subtle Border
        "highlight": "EAB308",    # Gold
    },
    "tech_charcoal": {
        "primary": "18181B",      # Zinc 900
        "secondary": "27272A",    # Zinc 800
        "accent": "6366F1",       # Indigo 500
        "accent_light": "EEF2FF", # Indigo 50
        "text": "18181B",         # Dark Zinc
        "text_muted": "71717A",   # Zinc 500
        "surface": "FAFAFA",      # Zinc 50
        "border": "D4D4D8",       # Zinc 300
        "border_subtle": "E4E4E7",# Zinc 200
        "highlight": "EC4899",    # Pink / Magenta
    },
}

FONT_FAMILIES = {
    "modern": {"display": "Segoe UI", "body": "Segoe UI", "code": "Consolas"},
    "executive": {"display": "Aptos Display", "body": "Aptos", "code": "Consolas"},
    "academic": {"display": "Georgia", "body": "Georgia", "code": "Courier New"},
    "clean": {"display": "Arial", "body": "Arial", "code": "Consolas"},
    "arabic": {"display": "Amiri", "body": "Amiri", "code": "Consolas"},
}

def hex_to_rgb(hex_code: str) -> RGBColor:
    """Convert hex string (e.g. '0F172A') to RGBColor."""
    hex_code = hex_code.lstrip("#")
    r = int(hex_code[0:2], 16)
    g = int(hex_code[2:4], 16)
    b = int(hex_code[4:6], 16)
    return RGBColor(r, g, b)

# ---------------------------------------------------------------------------
# 2. PAGE GEOMETRY & MARGINS
# ---------------------------------------------------------------------------

MARGIN_PRESETS = {
    "normal": {"top": 1.0, "bottom": 1.0, "left": 1.0, "right": 1.0, "header": 0.5, "footer": 0.5},
    "executive": {"top": 0.75, "bottom": 0.75, "left": 0.8, "right": 0.8, "header": 0.4, "footer": 0.4},
    "compact": {"top": 0.6, "bottom": 0.6, "left": 0.65, "right": 0.65, "header": 0.35, "footer": 0.35},
    "academic": {"top": 1.0, "bottom": 1.0, "left": 1.25, "right": 1.0, "header": 0.5, "footer": 0.5},
}

def apply_page_geometry(section, preset: str = "executive") -> None:
    """Apply consistent page margins and header/footer distances."""
    cfg = MARGIN_PRESETS.get(preset, MARGIN_PRESETS["executive"])
    section.top_margin = Inches(cfg["top"])
    section.bottom_margin = Inches(cfg["bottom"])
    section.left_margin = Inches(cfg["left"])
    section.right_margin = Inches(cfg["right"])
    section.header_distance = Inches(cfg["header"])
    section.footer_distance = Inches(cfg["footer"])

# ---------------------------------------------------------------------------
# 3. TYPOGRAPHY SETUP
# ---------------------------------------------------------------------------

def setup_document_styles(
    doc: Document,
    palette_name: str = "slate_executive",
    font_preset: str = "modern",
) -> Dict[str, str]:
    """Configure consistent base typography, scales, line spacing, and colors."""
    palette = PALETTES.get(palette_name, PALETTES["slate_executive"])
    fonts = FONT_FAMILIES.get(font_preset, FONT_FAMILIES["modern"])

    styles = doc.styles

    # Normal / Body Text
    normal = styles["Normal"]
    normal.font.name = fonts["body"]
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = hex_to_rgb(palette["text"])
    normal.paragraph_format.line_spacing = 1.2
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.space_before = Pt(0)

    # Heading 1
    if "Heading 1" in styles:
        h1 = styles["Heading 1"]
        h1.font.name = fonts["display"]
        h1.font.size = Pt(18)
        h1.font.bold = True
        h1.font.color.rgb = hex_to_rgb(palette["primary"])
        h1.paragraph_format.space_before = Pt(18)
        h1.paragraph_format.space_after = Pt(6)
        h1.paragraph_format.keep_with_next = True

    # Heading 2
    if "Heading 2" in styles:
        h2 = styles["Heading 2"]
        h2.font.name = fonts["display"]
        h2.font.size = Pt(14)
        h2.font.bold = True
        h2.font.color.rgb = hex_to_rgb(palette["secondary"])
        h2.paragraph_format.space_before = Pt(14)
        h2.paragraph_format.space_after = Pt(4)
        h2.paragraph_format.keep_with_next = True

    # Heading 3
    if "Heading 3" in styles:
        h3 = styles["Heading 3"]
        h3.font.name = fonts["display"]
        h3.font.size = Pt(12)
        h3.font.bold = True
        h3.font.color.rgb = hex_to_rgb(palette["accent"])
        h3.paragraph_format.space_before = Pt(10)
        h3.paragraph_format.space_after = Pt(3)
        h3.paragraph_format.keep_with_next = True

    return palette

def add_heading(
    doc: Document,
    text: str,
    level: int = 1,
    palette_name: str = "slate_executive",
    subtitle: Optional[str] = None,
) -> None:
    """
    Add a heading with strict keep_with_next to permanently prevent orphan headings.
    Optionally appends a small muted subtitle right underneath.
    """
    palette = PALETTES.get(palette_name, PALETTES["slate_executive"])
    p = doc.add_heading(text, level=level)
    p.paragraph_format.keep_with_next = True

    if subtitle:
        sub = doc.add_paragraph()
        sub.paragraph_format.space_before = Pt(0)
        sub.paragraph_format.space_after = Pt(8)
        sub.paragraph_format.keep_with_next = True
        run = sub.add_run(subtitle)
        run.font.size = Pt(9.5)
        run.font.italic = True
        run.font.color.rgb = hex_to_rgb(palette["text_muted"])

# ---------------------------------------------------------------------------
# 4. LOW-LEVEL OPENXML HELPERS (Shading, Margins, Borders, Bidi)
# ---------------------------------------------------------------------------

def set_cell_background(cell, hex_color: str) -> None:
    """Set background fill color of a table cell."""
    hex_color = hex_color.lstrip("#")
    tcPr = cell._tc.get_or_add_tcPr()
    for child in list(tcPr):
        if child.tag.endswith("shd"):
            tcPr.remove(child)
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{hex_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top: int = 120, bottom: int = 120, left: int = 160, right: int = 160) -> None:
    """
    Set internal padding for a cell in twips (1 pt = 20 twips).
    e.g., 120 twips = 6pt, 160 twips = 8pt.
    """
    tcPr = cell._tc.get_or_add_tcPr()
    for child in list(tcPr):
        if child.tag.endswith("tcMar"):
            tcPr.remove(child)
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>\n'
        f'  <w:top w:w="{top}" w:type="dxa"/>\n'
        f'  <w:bottom w:w="{bottom}" w:type="dxa"/>\n'
        f'  <w:left w:w="{left}" w:type="dxa"/>\n'
        f'  <w:right w:w="{right}" w:type="dxa"/>\n'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)

def set_table_cell_margins(table, top: int = 120, bottom: int = 120, left: int = 160, right: int = 160) -> None:
    """Set default padding for all cells in a table via tblCellMar."""
    tblPr = table._tbl.tblPr
    for child in list(tblPr):
        if child.tag.endswith("tblCellMar"):
            tblPr.remove(child)
    tblCellMar = parse_xml(
        f'<w:tblCellMar {nsdecls("w")}>\n'
        f'  <w:top w:w="{top}" w:type="dxa"/>\n'
        f'  <w:bottom w:w="{bottom}" w:type="dxa"/>\n'
        f'  <w:left w:w="{left}" w:type="dxa"/>\n'
        f'  <w:right w:w="{right}" w:type="dxa"/>\n'
        f'</w:tblCellMar>'
    )
    tblPr.append(tblCellMar)

def make_row_header(row) -> None:
    """Ensure table header row repeats across page breaks (w:tblHeader)."""
    trPr = row._tr.get_or_add_trPr()
    if trPr.find(qn("w:tblHeader")) is None:
        trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))

def prevent_row_split(row) -> None:
    """Prevent table row from splitting across a page break (w:cantSplit)."""
    trPr = row._tr.get_or_add_trPr()
    if trPr.find(qn("w:cantSplit")) is None:
        trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))

def set_callout_borders(cell, accent_hex: str, border_size: int = 36) -> None:
    """
    Set thick accent border on the left and none on top/right/bottom.
    border_size: 36 eighths of a pt = 4.5pt.
    """
    accent_hex = accent_hex.lstrip("#")
    tcPr = cell._tc.get_or_add_tcPr()
    for child in list(tcPr):
        if child.tag.endswith("tcBorders"):
            tcPr.remove(child)
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>\n'
        f'  <w:top w:val="none"/>\n'
        f'  <w:left w:val="single" w:sz="{border_size}" w:space="0" w:color="{accent_hex}"/>\n'
        f'  <w:bottom w:val="none"/>\n'
        f'  <w:right w:val="none"/>\n'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)

def set_table_horizontal_borders(table, border_color_hex: str = "E2E8F0", header_border_hex: str = "2563EB") -> None:
    """Apply clean, modern borders: subtle horizontal lines, no vertical borders."""
    border_color_hex = border_color_hex.lstrip("#")
    header_border_hex = header_border_hex.lstrip("#")
    tblPr = table._tbl.tblPr
    for child in list(tblPr):
        if child.tag.endswith("tblBorders"):
            tblPr.remove(child)
    tblBorders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>\n'
        f'  <w:top w:val="single" w:sz="6" w:space="0" w:color="{border_color_hex}"/>\n'
        f'  <w:bottom w:val="single" w:sz="12" w:space="0" w:color="{header_border_hex}"/>\n'
        f'  <w:left w:val="none"/>\n'
        f'  <w:right w:val="none"/>\n'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{border_color_hex}"/>\n'
        f'  <w:insideV w:val="none"/>\n'
        f'</w:tblBorders>'
    )
    tblPr.append(tblBorders)

def set_arabic_rtl(element, is_table: bool = False) -> None:
    """Enable RTL for a paragraph, run, or table."""
    if is_table:
        tblPr = element._tbl.tblPr
        if tblPr.find(qn("w:bidiVisual")) is None:
            tblPr.append(parse_xml(f'<w:bidiVisual {nsdecls("w")}/>'))
    else:
        pPr = element._p.get_or_add_pPr()
        if pPr.find(qn("w:bidi")) is None:
            pPr.append(parse_xml(f'<w:bidi {nsdecls("w")}/>'))

# ---------------------------------------------------------------------------
# 5. HIGH-LEVEL COMPONENT BUILDERS (Tables, Callouts, Cards)
# ---------------------------------------------------------------------------

def create_styled_table(
    doc: Document,
    headers: List[str],
    rows: List[List[Union[str, int, float]]],
    col_widths: Optional[List[float]] = None,
    palette_name: str = "slate_executive",
    zebra_striping: bool = True,
    alignment: int = WD_TABLE_ALIGNMENT.CENTER,
) -> docx.table.Table:
    """
    Creates a publication-ready, scannable table with:
    - Auto-repeated header row across pages (w:tblHeader)
    - Non-splitting rows (w:cantSplit)
    - Generous cell padding (top/bottom 6pt, left/right 8pt)
    - Subtle borders & clean typography
    - Optional zebra striping
    """
    palette = PALETTES.get(palette_name, PALETTES["slate_executive"])
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.alignment = alignment
    table.autofit = False

    # Default cell margins across table
    set_table_cell_margins(table, top=140, bottom=140, left=180, right=180)
    set_table_horizontal_borders(table, border_color_hex=palette["border_subtle"], header_border_hex=palette["accent"])

    # Header Row
    header_row = table.rows[0]
    make_row_header(header_row)
    prevent_row_split(header_row)

    for i, title in enumerate(headers):
        cell = header_row.cells[i]
        set_cell_background(cell, palette["primary"])
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(str(title))
        run.font.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = hex_to_rgb("FFFFFF")

    # Data Rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        prevent_row_split(row)

        bg_color = palette["surface"] if (zebra_striping and r_idx % 2 == 1) else "FFFFFF"

        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            if bg_color != "FFFFFF":
                set_cell_background(cell, bg_color)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(str(val))
            run.font.size = Pt(9.5)
            run.font.color.rgb = hex_to_rgb(palette["text"])

    if col_widths:
        for r in table.rows:
            for c_idx, width in enumerate(col_widths):
                if c_idx < len(r.cells):
                    r.cells[c_idx].width = Inches(width)

    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_before = Pt(4)
    p_after.paragraph_format.space_after = Pt(8)

    return table

def add_callout(
    doc: Document,
    text: str,
    title: Optional[str] = None,
    kind: str = "takeaway",
    palette_name: str = "slate_executive",
) -> docx.table.Table:
    """
    Creates an elegant, high-impact callout box:
    - 1x1 table with thick accent left border
    - Soft tint background
    - Internal padding
    - Keep with next to prevent orphan boxes
    """
    palette = PALETTES.get(palette_name, PALETTES["slate_executive"])

    color_map = {
        "takeaway": (palette["accent"], palette["accent_light"]),
        "note": (palette["secondary"], palette["surface"]),
        "warning": (palette.get("highlight", "F59E0B"), "FFFBEB"),
        "success": ("10B981", "ECFDF5"),
    }
    accent_hex, bg_hex = color_map.get(kind, (palette["accent"], palette["accent_light"]))

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    cell = table.cell(0, 0)
    cell.width = Inches(6.8)

    set_cell_background(cell, bg_hex)
    set_callout_borders(cell, accent_hex=accent_hex, border_size=36)
    set_cell_margins(cell, top=160, bottom=160, left=200, right=180)
    prevent_row_split(table.rows[0])

    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3 if title else 0)
    p.paragraph_format.line_spacing = 1.15

    if title:
        run_title = p.add_run(title.upper() + "\n")
        run_title.font.bold = True
        run_title.font.size = Pt(8.5)
        run_title.font.color.rgb = hex_to_rgb(accent_hex)

    run_text = p.add_run(text)
    run_text.font.size = Pt(9.5)
    run_text.font.color.rgb = hex_to_rgb(palette["text"])

    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_before = Pt(2)
    p_after.paragraph_format.space_after = Pt(6)

    return table

def add_stat_card_row(
    doc: Document,
    stats: List[Tuple[str, str, Optional[str]]],
    palette_name: str = "slate_executive",
) -> docx.table.Table:
    """
    Creates an executive metric summary row (e.g. 3 or 4 stat cards):
    stats = [(Value, Label, Subtitle/Change), ...]
    """
    palette = PALETTES.get(palette_name, PALETTES["slate_executive"])
    cols = len(stats)
    table = doc.add_table(rows=1, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    set_table_cell_margins(table, top=140, bottom=140, left=160, right=160)

    card_width = Inches(6.8 / cols)

    for i, (val, label, sub) in enumerate(stats):
        cell = table.cell(0, i)
        cell.width = card_width
        set_cell_background(cell, palette["surface"])

        tcPr = cell._tc.get_or_add_tcPr()
        borders = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>\n'
            f'  <w:top w:val="single" w:sz="6" w:space="0" w:color="{palette["border_subtle"]}"/>\n'
            f'  <w:bottom w:val="single" w:sz="18" w:space="0" w:color="{palette["accent"]}"/>\n'
            f'  <w:left w:val="single" w:sz="6" w:space="0" w:color="{palette["border_subtle"]}"/>\n'
            f'  <w:right w:val="single" w:sz="6" w:space="0" w:color="{palette["border_subtle"]}"/>\n'
            f'</w:tcBorders>'
        )
        tcPr.append(borders)

        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)

        run_val = p.add_run(f"{val}\n")
        run_val.font.bold = True
        run_val.font.size = Pt(18)
        run_val.font.color.rgb = hex_to_rgb(palette["primary"])

        run_label = p.add_run(label)
        run_label.font.size = Pt(8.5)
        run_label.font.bold = True
        run_label.font.color.rgb = hex_to_rgb(palette["text_muted"])

        if sub:
            p.add_run(f"\n{sub}").font.size = Pt(8)

    prevent_row_split(table.rows[0])
    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_after = Pt(8)
    return table

# ---------------------------------------------------------------------------
# 6. RUNNING HEADERS, FOOTERS & NATIVE PAGE NUMBERS
# ---------------------------------------------------------------------------

def setup_header_footer(
    doc: Document,
    title_left: str,
    doc_info_right: str = "",
    palette_name: str = "slate_executive",
    show_page_numbers: bool = True,
    suppress_cover_page: bool = True,
) -> None:
    """
    Configures clean running header and footer:
    - Suppresses header/footer on cover page (different_first_page_header_footer)
    - Top header: document title (left) + category/doc info (right) + subtle divider
    - Bottom footer: copyright/institution (left) + Page X of Y (right)
    """
    palette = PALETTES.get(palette_name, PALETTES["slate_executive"])
    section = doc.sections[0]
    if suppress_cover_page:
        section.different_first_page_header_footer = True

    # Running Header
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_before = Pt(0)
    hp.paragraph_format.space_after = Pt(4)

    run_hl = hp.add_run(title_left)
    run_hl.font.size = Pt(8.5)
    run_hl.font.color.rgb = hex_to_rgb(palette["text_muted"])

    if doc_info_right:
        run_right = hp.add_run(f"  |  {doc_info_right}")
        run_right.font.size = Pt(8.5)
        run_right.font.color.rgb = hex_to_rgb(palette["text_muted"])

    pPr = hp._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>\n'
        f'  <w:bottom w:val="single" w:sz="4" w:space="4" w:color="{palette["border_subtle"]}"/>\n'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)

    # Running Footer
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(4)
    fp.paragraph_format.space_after = Pt(0)

    fpPr = fp._p.get_or_add_pPr()
    fpBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>\n'
        f'  <w:top w:val="single" w:sz="4" w:space="4" w:color="{palette["border_subtle"]}"/>\n'
        f'</w:pBdr>'
    )
    fpPr.append(fpBdr)

    if show_page_numbers:
        r_txt = fp.add_run("Page ")
        r_txt.font.size = Pt(8.5)
        r_txt.font.color.rgb = hex_to_rgb(palette["text_muted"])

        fld_page = parse_xml(
            f'<w:fldSimple {nsdecls("w")} w:instr="PAGE">\n'
            f'  <w:r><w:rPr><w:sz w:val="17"/><w:color w:val="{palette["text_muted"]}"/></w:rPr><w:t>1</w:t></w:r>\n'
            f'</w:fldSimple>'
        )
        fp._p.append(fld_page)

        r_of = fp.add_run(" of ")
        r_of.font.size = Pt(8.5)
        r_of.font.color.rgb = hex_to_rgb(palette["text_muted"])

        fld_numpages = parse_xml(
            f'<w:fldSimple {nsdecls("w")} w:instr="NUMPAGES">\n'
            f'  <w:r><w:rPr><w:sz w:val="17"/><w:color w:val="{palette["text_muted"]}"/></w:rPr><w:t>1</w:t></w:r>\n'
            f'</w:fldSimple>'
        )
        fp._p.append(fld_numpages)

# ---------------------------------------------------------------------------
# 7. COVER PAGE ARCHETYPES
# ---------------------------------------------------------------------------

def add_cover_page(
    doc: Document,
    title: str,
    subtitle: str,
    author_org: str,
    date_version: str,
    archetype: str = "executive_stripe",
    palette_name: str = "slate_executive",
) -> None:
    """
    Creates a standout cover page with proper whitespace, hierarchy, and page break.
    Archetypes:
    - 'executive_stripe': Accent color block, bold display typography, clean metadata table.
    - 'academic_clean': Formal, restrained, centered institutional block.
    - 'minimal_modern': Asymmetric, high whitespace, bold typography.
    """
    palette = PALETTES.get(palette_name, PALETTES["slate_executive"])

    if archetype == "executive_stripe":
        p_top = doc.add_paragraph()
        p_top.paragraph_format.space_before = Pt(40)

        stripe = doc.add_table(rows=1, cols=1)
        stripe.autofit = False
        stripe.cell(0, 0).width = Inches(6.8)
        set_cell_background(stripe.cell(0, 0), palette["primary"])
        set_cell_margins(stripe.cell(0, 0), top=80, bottom=80, left=140, right=140)
        p_str = stripe.cell(0, 0).paragraphs[0]
        r_str = p_str.add_run("OFFICIAL REPORT  |  CONFIDENTIAL")
        r_str.font.size = Pt(8.5)
        r_str.font.bold = True
        r_str.font.color.rgb = hex_to_rgb(palette["accent_light"])

        p_title = doc.add_paragraph()
        p_title.paragraph_format.space_before = Pt(36)
        p_title.paragraph_format.space_after = Pt(12)
        p_title.paragraph_format.keep_with_next = True
        r_title = p_title.add_run(title)
        r_title.font.bold = True
        r_title.font.size = Pt(30)
        r_title.font.color.rgb = hex_to_rgb(palette["primary"])

        p_sub = doc.add_paragraph()
        p_sub.paragraph_format.space_before = Pt(0)
        p_sub.paragraph_format.space_after = Pt(48)
        p_sub.paragraph_format.keep_with_next = True
        r_sub = p_sub.add_run(subtitle)
        r_sub.font.size = Pt(13)
        r_sub.font.color.rgb = hex_to_rgb(palette["text_muted"])

        p_div = doc.add_paragraph()
        p_div.paragraph_format.space_after = Pt(60)
        pPr = p_div._p.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>\n'
            f'  <w:bottom w:val="single" w:sz="18" w:space="1" w:color="{palette["accent"]}"/>\n'
            f'</w:pBdr>'
        )
        pPr.append(pBdr)

        p_meta = doc.add_paragraph()
        p_meta.paragraph_format.space_before = Pt(80)
        p_meta.paragraph_format.space_after = Pt(2)
        r_org = p_meta.add_run(f"PREPARED BY\n{author_org}\n\n")
        r_org.font.size = Pt(9.5)
        r_org.font.color.rgb = hex_to_rgb(palette["text_muted"])

        r_dt = p_meta.add_run(f"DATE & VERSION\n{date_version}")
        r_dt.font.size = Pt(9.5)
        r_dt.font.bold = True
        r_dt.font.color.rgb = hex_to_rgb(palette["primary"])

    elif archetype == "academic_clean":
        p_top = doc.add_paragraph()
        p_top.paragraph_format.space_before = Pt(80)

        p_inst = doc.add_paragraph()
        p_inst.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_inst.paragraph_format.space_after = Pt(36)
        r_inst = p_inst.add_run(author_org.upper())
        r_inst.font.size = Pt(11)
        r_inst.font.bold = True
        r_inst.font.color.rgb = hex_to_rgb(palette["text_muted"])

        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_title.paragraph_format.space_before = Pt(40)
        p_title.paragraph_format.space_after = Pt(16)
        r_title = p_title.add_run(title)
        r_title.font.bold = True
        r_title.font.size = Pt(26)
        r_title.font.color.rgb = hex_to_rgb(palette["primary"])

        p_sub = doc.add_paragraph()
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_sub.paragraph_format.space_after = Pt(120)
        r_sub = p_sub.add_run(subtitle)
        r_sub.font.size = Pt(12)
        r_sub.font.italic = True
        r_sub.font.color.rgb = hex_to_rgb(palette["text_muted"])

        p_date = doc.add_paragraph()
        p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_date.paragraph_format.space_before = Pt(60)
        r_date = p_date.add_run(date_version)
        r_date.font.size = Pt(10)
        r_date.font.color.rgb = hex_to_rgb(palette["text"])

    else:  # minimal_modern
        p_top = doc.add_paragraph()
        p_top.paragraph_format.space_before = Pt(60)

        p_badge = doc.add_paragraph()
        p_badge.paragraph_format.space_after = Pt(20)
        r_badge = p_badge.add_run("RESEARCH & STRATEGY BRIEF")
        r_badge.font.size = Pt(9)
        r_badge.font.bold = True
        r_badge.font.color.rgb = hex_to_rgb(palette["accent"])

        p_title = doc.add_paragraph()
        p_title.paragraph_format.space_after = Pt(14)
        r_title = p_title.add_run(title)
        r_title.font.bold = True
        r_title.font.size = Pt(28)
        r_title.font.color.rgb = hex_to_rgb(palette["primary"])

        p_sub = doc.add_paragraph()
        p_sub.paragraph_format.space_after = Pt(80)
        r_sub = p_sub.add_run(subtitle)
        r_sub.font.size = Pt(12)
        r_sub.font.color.rgb = hex_to_rgb(palette["text_muted"])

        p_info = doc.add_paragraph()
        p_info.paragraph_format.space_before = Pt(100)
        r_info = p_info.add_run(f"{author_org}  ·  {date_version}")
        r_info.font.size = Pt(9.5)
        r_info.font.color.rgb = hex_to_rgb(palette["text_muted"])

    doc.add_page_break()

# ---------------------------------------------------------------------------
# 8. HEADLESS RENDERING & VERIFICATION (LibreOffice + pdftoppm)
# ---------------------------------------------------------------------------

def render_to_pdf(docx_path: Union[str, Path], output_dir: Optional[Union[str, Path]] = None) -> Path:
    """
    Render DOCX to PDF using headless LibreOffice.
    Returns path to the generated PDF.
    """
    docx_file = Path(docx_path).resolve()
    if not docx_file.exists():
        raise FileNotFoundError(f"DOCX file not found: {docx_file}")

    out_dir = Path(output_dir).resolve() if output_dir else docx_file.parent
    out_dir.mkdir(parents=True, exist_ok=True)

    cmd = ["soffice", "--headless", "--convert-to", "pdf", str(docx_file), "--outdir", str(out_dir)]
    res = subprocess.run(cmd, capture_output=True, text=True, check=False)
    if res.returncode != 0:
        raise RuntimeError(f"LibreOffice PDF conversion failed:\nSTDOUT: {res.stdout}\nSTDERR: {res.stderr}")

    expected_pdf = out_dir / (docx_file.stem + ".pdf")
    if not expected_pdf.exists():
        raise FileNotFoundError(f"Expected PDF output was not created: {expected_pdf}")
    return expected_pdf

def render_pdf_to_images(
    pdf_path: Union[str, Path],
    output_dir: Optional[Union[str, Path]] = None,
    dpi: int = 150,
    prefix: str = "page",
) -> List[Path]:
    """
    Render each PDF page as a high-resolution PNG using pdftoppm.
    Returns sorted list of generated PNG image paths.
    """
    pdf_file = Path(pdf_path).resolve()
    if not pdf_file.exists():
        raise FileNotFoundError(f"PDF file not found: {pdf_file}")

    out_dir = Path(output_dir).resolve() if output_dir else pdf_file.parent / "pages"
    out_dir.mkdir(parents=True, exist_ok=True)

    output_prefix = out_dir / prefix
    cmd = ["pdftoppm", "-png", "-r", str(dpi), str(pdf_file), str(output_prefix)]
    res = subprocess.run(cmd, capture_output=True, text=True, check=False)
    if res.returncode != 0:
        raise RuntimeError(f"pdftoppm image extraction failed:\nSTDOUT: {res.stdout}\nSTDERR: {res.stderr}")

    images = sorted(list(out_dir.glob(f"{prefix}-*.png")))
    return images

def render_and_preview(docx_path: Union[str, Path], output_dir: Optional[Union[str, Path]] = None) -> Tuple[Path, List[Path]]:
    """Convenience pipeline: DOCX -> PDF -> PNG page images."""
    pdf = render_to_pdf(docx_path, output_dir=output_dir)
    images = render_pdf_to_images(pdf, output_dir=output_dir)
    return pdf, images
