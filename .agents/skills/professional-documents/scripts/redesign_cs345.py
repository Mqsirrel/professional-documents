"""
redesign_cs345.py - High-Taste Redesign of CS345 Operating Systems Lab 1 Template

Applies K3-informed professional document architecture:
- Authoritative academic/technical art direction
- Official Taibah University logo and Arabic header integration
- Clean typography (Segoe UI, Consolas, Arial/Traditional Arabic)
- Distinct command badges and terminal answer boxes
- Padded screenshot drop containers
- Anti-orphan and anti-split table discipline
- Running headers, footers, and dynamic Word page numbers (Page X of Y)
"""

import os
import subprocess
from pathlib import Path
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn

# Curated Palette: Slate Executive / Academic Navy
COLOR_PRIMARY = "0F172A"       # Slate 900 (Deep Navy)
COLOR_SECONDARY = "1E293B"     # Slate 800
COLOR_ACCENT = "2563EB"        # Royal Blue
COLOR_ACCENT_LIGHT = "EFF6FF"  # Blue 50
COLOR_TEXT = "0F172A"          # Dark text
COLOR_TEXT_MUTED = "64748B"    # Slate 500
COLOR_SURFACE = "F8FAFC"       # Slate 50
COLOR_SURFACE_ALT = "F1F5F9"   # Slate 100
COLOR_BORDER = "CBD5E1"        # Slate 300
COLOR_BORDER_LIGHT = "E2E8F0"  # Slate 200

def hex_to_rgb(hex_code: str) -> RGBColor:
    hex_code = hex_code.lstrip("#")
    return RGBColor(int(hex_code[0:2], 16), int(hex_code[2:4], 16), int(hex_code[4:6], 16))

def set_cell_shading(cell, hex_color: str) -> None:
    hex_color = hex_color.lstrip("#")
    tcPr = cell._tc.get_or_add_tcPr()
    for child in list(tcPr):
        if child.tag.endswith("shd"):
            tcPr.remove(child)
    tcPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{hex_color}"/>'))

def set_cell_padding(cell, top: int = 120, bottom: int = 120, left: int = 160, right: int = 160) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    for child in list(tcPr):
        if child.tag.endswith("tcMar"):
            tcPr.remove(child)
    tcPr.append(parse_xml(
        f'<w:tcMar {nsdecls("w")}>\n'
        f'  <w:top w:w="{top}" w:type="dxa"/>\n'
        f'  <w:bottom w:w="{bottom}" w:type="dxa"/>\n'
        f'  <w:left w:w="{left}" w:type="dxa"/>\n'
        f'  <w:right w:w="{right}" w:type="dxa"/>\n'
        f'</w:tcMar>'
    ))

def set_table_padding(table, top: int = 120, bottom: int = 120, left: int = 160, right: int = 160) -> None:
    tblPr = table._tbl.tblPr
    for child in list(tblPr):
        if child.tag.endswith("tblCellMar"):
            tblPr.remove(child)
    tblPr.append(parse_xml(
        f'<w:tblCellMar {nsdecls("w")}>\n'
        f'  <w:top w:w="{top}" w:type="dxa"/>\n'
        f'  <w:bottom w:w="{bottom}" w:type="dxa"/>\n'
        f'  <w:left w:w="{left}" w:type="dxa"/>\n'
        f'  <w:right w:w="{right}" w:type="dxa"/>\n'
        f'</w:tblCellMar>'
    ))

def make_row_cant_split(row) -> None:
    trPr = row._tr.get_or_add_trPr()
    if trPr.find(qn("w:cantSplit")) is None:
        trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))

def set_bidi(p) -> None:
    pPr = p._p.get_or_add_pPr()
    if pPr.find(qn("w:bidi")) is None:
        pPr.append(parse_xml(f'<w:bidi {nsdecls("w")}/>'))

def set_arabic_font(run, font_name="Sakkal Majalla", sz=None, bold=None, color_hex=None):
    rPr = run._r.get_or_add_rPr()
    for c in list(rPr):
        if c.tag.endswith("rFonts"):
            rPr.remove(c)
    rPr.append(parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{font_name}" w:hAnsi="{font_name}" w:cs="{font_name}"/>'))
    if sz is not None:
        run.font.size = Pt(sz)
    if bold is not None:
        run.font.bold = bold
    if color_hex is not None:
        run.font.color.rgb = hex_to_rgb(color_hex)
    if rPr.find(qn("w:rtl")) is None:
        rPr.append(parse_xml(f'<w:rtl {nsdecls("w")}/>'))

def build_redesigned_document(output_path: str, logo_path: str) -> Document:
    doc = Document()

    # 1. Page Margins (Executive: 0.75" top/bottom, 0.8" left/right)
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.header_distance = Inches(0.4)
    section.footer_distance = Inches(0.4)
    section.different_first_page_header_footer = True

    # 2. Typography Styles
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Segoe UI"
    normal.font.size = Pt(10)
    normal.font.color.rgb = hex_to_rgb(COLOR_TEXT)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.space_before = Pt(0)
    nrPr = normal._element.get_or_add_rPr()
    nrPr.append(parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="Segoe UI" w:hAnsi="Segoe UI" w:cs="Sakkal Majalla"/>'))

    # -----------------------------------------------------------------------
    # COVER PAGE
    # -----------------------------------------------------------------------
    
    # Top Header Table: Logo (Left) + Arabic Institutional Title (Right)
    header_tbl = doc.add_table(rows=1, cols=2)
    header_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_tbl.autofit = False
    set_table_padding(header_tbl, top=40, bottom=40, left=60, right=60)
    
    # Cell 0: Logo
    cell_logo = header_tbl.cell(0, 0)
    cell_logo.width = Inches(2.8)
    p_logo = cell_logo.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_logo.paragraph_format.space_before = Pt(0)
    p_logo.paragraph_format.space_after = Pt(0)
    if os.path.exists(logo_path):
        p_logo.add_run().add_picture(logo_path, width=Inches(1.85))

    # Cell 1: Arabic Institutional Info
    cell_ar = header_tbl.cell(0, 1)
    cell_ar.width = Inches(4.0)
    p_ar = cell_ar.paragraphs[0]
    p_ar.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_bidi(p_ar)
    p_ar.paragraph_format.space_before = Pt(0)
    p_ar.paragraph_format.space_after = Pt(0)
    p_ar.paragraph_format.line_spacing = 1.25

    ar_lines = [
        ("المملكة العربية السعودية", 9.5, True),
        ("وزارة التعليم العالي", 9.0, False),
        ("جامعة طيبة", 9.5, True),
        ("كلية علوم وهندسة الحاسب الآلي", 9.0, False),
        ("قسم علوم الحاسب", 9.0, True),
    ]
    for idx, (line_text, sz, is_bld) in enumerate(ar_lines):
        r = p_ar.add_run(line_text + ("\n" if idx < len(ar_lines)-1 else ""))
        set_arabic_font(r, font_name="Sakkal Majalla", sz=sz + 2.0, bold=is_bld, color_hex=COLOR_PRIMARY)

    make_row_cant_split(header_tbl.rows[0])

    # Decorative Accent Divider Line
    p_rule = doc.add_paragraph()
    p_rule.paragraph_format.space_before = Pt(16)
    p_rule.paragraph_format.space_after = Pt(36)
    pPr = p_rule._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>\n'
        f'  <w:bottom w:val="single" w:sz="18" w:space="1" w:color="{COLOR_ACCENT}"/>\n'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)

    # Course Eyebrow Tag
    p_tag = doc.add_paragraph()
    p_tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_tag.paragraph_format.space_after = Pt(6)
    p_tag.paragraph_format.keep_with_next = True
    r_tag = p_tag.add_run("CS345  •  OPERATING SYSTEMS LAB")
    r_tag.font.size = Pt(10.5)
    r_tag.font.bold = True
    r_tag.font.color.rgb = hex_to_rgb(COLOR_ACCENT)

    # Document Main Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(12)
    p_title.paragraph_format.keep_with_next = True
    r_title = p_title.add_run("Operating Systems Lab 1")
    r_title.font.size = Pt(30)
    r_title.font.bold = True
    r_title.font.color.rgb = hex_to_rgb(COLOR_PRIMARY)

    # Subtitle
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_before = Pt(0)
    p_sub.paragraph_format.space_after = Pt(40)
    p_sub.paragraph_format.keep_with_next = True
    r_sub = p_sub.add_run("Lab Report: Essential Linux Commands & Process Management")
    r_sub.font.size = Pt(13)
    r_sub.font.color.rgb = hex_to_rgb(COLOR_TEXT_MUTED)

    # Submission Card Table (Students Team)
    sub_tbl = doc.add_table(rows=5, cols=2)
    sub_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    sub_tbl.autofit = False
    set_table_padding(sub_tbl, top=130, bottom=130, left=180, right=180)

    # Card borders: subtle box with accent top rule
    tblPr = sub_tbl._tbl.tblPr
    tblBorders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>\n'
        f'  <w:top w:val="single" w:sz="18" w:space="0" w:color="{COLOR_ACCENT}"/>\n'
        f'  <w:bottom w:val="single" w:sz="6" w:space="0" w:color="{COLOR_BORDER}"/>\n'
        f'  <w:left w:val="single" w:sz="6" w:space="0" w:color="{COLOR_BORDER_LIGHT}"/>\n'
        f'  <w:right w:val="single" w:sz="6" w:space="0" w:color="{COLOR_BORDER_LIGHT}"/>\n'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{COLOR_BORDER_LIGHT}"/>\n'
        f'  <w:insideV w:val="none"/>\n'
        f'</w:tblBorders>'
    )
    tblPr.append(tblBorders)

    # Header Row (Merged across both columns)
    hdr_cell = sub_tbl.cell(0, 0)
    sub_tbl.cell(0, 1).merge(hdr_cell)
    set_cell_shading(hdr_cell, COLOR_SURFACE_ALT)
    p_shdr = hdr_cell.paragraphs[0]
    p_shdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_shdr.paragraph_format.space_before = Pt(2)
    p_shdr.paragraph_format.space_after = Pt(2)
    r_shdr = p_shdr.add_run("SUBMITTED BY  /  إعداد الطلاب")
    r_shdr.font.size = Pt(9.5)
    r_shdr.font.bold = True
    r_shdr.font.color.rgb = hex_to_rgb(COLOR_PRIMARY)

    students = [
        ("البراء فهد مشعان موقد", "ID: 4504099"),
        ("فائز فيصل الحربي", "ID: 4500459"),
        ("الياس عيد علي المحمدي", "ID: 4501089"),
    ]

    for idx, (name, s_id) in enumerate(students):
        row = sub_tbl.rows[idx + 1]
        make_row_cant_split(row)
        c_name = row.cells[0]
        c_id = row.cells[1]
        c_name.width = Inches(4.2)
        c_id.width = Inches(2.6)

        # Alternating subtle tint
        if idx % 2 == 1:
            set_cell_shading(c_name, COLOR_SURFACE)
            set_cell_shading(c_id, COLOR_SURFACE)

        # Student Name (RTL right-aligned)
        pn = c_name.paragraphs[0]
        pn.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_bidi(pn)
        pn.paragraph_format.space_before = Pt(1)
        pn.paragraph_format.space_after = Pt(1)
        rn = pn.add_run(name)
        set_arabic_font(rn, font_name="Sakkal Majalla", sz=12.5, bold=True, color_hex=COLOR_PRIMARY)

        # Student ID (Mono, left-aligned)
        pi = c_id.paragraphs[0]
        pi.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pi.paragraph_format.space_before = Pt(1)
        pi.paragraph_format.space_after = Pt(1)
        ri = pi.add_run(s_id)
        ri.font.name = "Consolas"
        ri.font.bold = True
        ri.font.size = Pt(10)
        ri.font.color.rgb = hex_to_rgb(COLOR_SECONDARY)

    # Footer Row (Metadata)
    ftr_cell = sub_tbl.cell(4, 0)
    sub_tbl.cell(4, 1).merge(ftr_cell)
    set_cell_shading(ftr_cell, COLOR_SURFACE_ALT)
    p_sftr = ftr_cell.paragraphs[0]
    p_sftr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sftr.paragraph_format.space_before = Pt(3)
    p_sftr.paragraph_format.space_after = Pt(3)
    r_sftr = p_sftr.add_run("COURSE: CS345  •  SEMESTER: Term 2, 1447H  •  DATE: September 2026")
    r_sftr.font.size = Pt(8.5)
    r_sftr.font.bold = True
    r_sftr.font.color.rgb = hex_to_rgb(COLOR_TEXT_MUTED)

    # Cover Bottom Institutional Text
    p_cbot = doc.add_paragraph()
    p_cbot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cbot.paragraph_format.space_before = Pt(70)
    p_cbot.paragraph_format.space_after = Pt(0)
    r_cbot = p_cbot.add_run("College of Computer Science & Engineering  •  Taibah University, Madinah")
    r_cbot.font.size = Pt(9.5)
    r_cbot.font.color.rgb = hex_to_rgb(COLOR_TEXT_MUTED)

    # Page Break after Cover Page
    doc.add_page_break()

    # -----------------------------------------------------------------------
    # RUNNING HEADERS & FOOTERS (From Page 2 Onwards)
    # -----------------------------------------------------------------------
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_before = Pt(0)
    hp.paragraph_format.space_after = Pt(4)

    r_hl = hp.add_run("CS345: Operating Systems Lab 1")
    r_hl.font.size = Pt(8.5)
    r_hl.font.bold = True
    r_hl.font.color.rgb = hex_to_rgb(COLOR_TEXT_MUTED)

    r_hr = hp.add_run("  |  Essential Linux Commands & Process Management")
    r_hr.font.size = Pt(8.5)
    r_hr.font.color.rgb = hex_to_rgb(COLOR_TEXT_MUTED)

    hpPr = hp._p.get_or_add_pPr()
    hpBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>\n'
        f'  <w:bottom w:val="single" w:sz="4" w:space="4" w:color="{COLOR_BORDER_LIGHT}"/>\n'
        f'</w:pBdr>'
    )
    hpPr.append(hpBdr)

    # Running Footer: College Info (Left) + Page X of Y (Right)
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(4)
    fp.paragraph_format.space_after = Pt(0)

    fpPr = fp._p.get_or_add_pPr()
    fpBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>\n'
        f'  <w:top w:val="single" w:sz="4" w:space="4" w:color="{COLOR_BORDER_LIGHT}"/>\n'
        f'</w:pBdr>'
    )
    fpPr.append(fpBdr)

    r_fl = fp.add_run("College of Computer Science & Engineering  •  Taibah University            ")
    r_fl.font.size = Pt(8.5)
    r_fl.font.color.rgb = hex_to_rgb(COLOR_TEXT_MUTED)

    r_pg_txt = fp.add_run("Page ")
    r_pg_txt.font.size = Pt(8.5)
    r_pg_txt.font.color.rgb = hex_to_rgb(COLOR_TEXT_MUTED)

    fld_page = parse_xml(
        f'<w:fldSimple {nsdecls("w")} w:instr="PAGE">\n'
        f'  <w:r><w:rPr><w:sz w:val="17"/><w:color w:val="{COLOR_TEXT_MUTED}"/></w:rPr><w:t>2</w:t></w:r>\n'
        f'</w:fldSimple>'
    )
    fp._p.append(fld_page)

    r_of = fp.add_run(" of ")
    r_of.font.size = Pt(8.5)
    r_of.font.color.rgb = hex_to_rgb(COLOR_TEXT_MUTED)

    fld_numpages = parse_xml(
        f'<w:fldSimple {nsdecls("w")} w:instr="NUMPAGES">\n'
        f'  <w:r><w:rPr><w:sz w:val="17"/><w:color w:val="{COLOR_TEXT_MUTED}"/></w:rPr><w:t>6</w:t></w:r>\n'
        f'</w:fldSimple>'
    )
    fp._p.append(fld_numpages)

    # -----------------------------------------------------------------------
    # INSIDE CONTENT: LAB INSTRUCTIONS & TASKS
    # -----------------------------------------------------------------------

    # Document Header on Page 2
    p_maintitle = doc.add_paragraph()
    p_maintitle.paragraph_format.space_before = Pt(8)
    p_maintitle.paragraph_format.space_after = Pt(2)
    p_maintitle.paragraph_format.keep_with_next = True
    r_mt = p_maintitle.add_run("Lab Tasks & Execution Verification")
    r_mt.font.size = Pt(18)
    r_mt.font.bold = True
    r_mt.font.color.rgb = hex_to_rgb(COLOR_PRIMARY)

    p_mainsub = doc.add_paragraph()
    p_mainsub.paragraph_format.space_before = Pt(0)
    p_mainsub.paragraph_format.space_after = Pt(10)
    p_mainsub.paragraph_format.keep_with_next = True
    r_ms = p_mainsub.add_run("Complete all 13 questions across 5 core sections. Include command syntax and execution screenshot.")
    r_ms.font.size = Pt(10)
    r_ms.font.color.rgb = hex_to_rgb(COLOR_TEXT_MUTED)

    # Instructions Callout Box
    inst_tbl = doc.add_table(rows=1, cols=1)
    inst_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    inst_tbl.autofit = False
    c_inst = inst_tbl.cell(0, 0)
    c_inst.width = Inches(6.8)
    set_cell_shading(c_inst, COLOR_ACCENT_LIGHT)
    set_cell_padding(c_inst, top=140, bottom=140, left=180, right=160)

    # Left accent border
    tcPr = c_inst._tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>\n'
        f'  <w:top w:val="none"/>\n'
        f'  <w:left w:val="single" w:sz="36" w:space="0" w:color="{COLOR_ACCENT}"/>\n'
        f'  <w:bottom w:val="none"/>\n'
        f'  <w:right w:val="none"/>\n'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)
    make_row_cant_split(inst_tbl.rows[0])

    p_ic = c_inst.paragraphs[0]
    p_ic.paragraph_format.space_before = Pt(0)
    p_ic.paragraph_format.space_after = Pt(0)
    p_ic.paragraph_format.line_spacing = 1.15
    r_ititle = p_ic.add_run("INSTRUCTIONS & SUBMISSION GUIDELINES\n")
    r_ititle.font.bold = True
    r_ititle.font.size = Pt(8.5)
    r_ititle.font.color.rgb = hex_to_rgb(COLOR_ACCENT)

    r_ibody = p_ic.add_run(
        "Execute each task in your Linux terminal. For each question, type the exact command used in the "
        "Command area and paste a clear terminal screenshot verifying execution and standard output in the dedicated box."
    )
    r_ibody.font.size = Pt(9.5)
    r_ibody.font.color.rgb = hex_to_rgb(COLOR_TEXT)

    p_sp = doc.add_paragraph()
    p_sp.paragraph_format.space_after = Pt(8)

    # -----------------------------------------------------------------------
    # QUESTION SECTIONS HELPER
    # -----------------------------------------------------------------------

    def add_section_header(sec_title: str, commands: list[str]) -> None:
        p_sec = doc.add_paragraph()
        p_sec.paragraph_format.space_before = Pt(14)
        p_sec.paragraph_format.space_after = Pt(4)
        p_sec.paragraph_format.keep_with_next = True

        r_sname = p_sec.add_run(sec_title)
        r_sname.font.size = Pt(12.5)
        r_sname.font.bold = True
        r_sname.font.color.rgb = hex_to_rgb(COLOR_PRIMARY)

        # Commands pill row
        p_cmd = doc.add_paragraph()
        p_cmd.paragraph_format.space_before = Pt(0)
        p_cmd.paragraph_format.space_after = Pt(10)
        p_cmd.paragraph_format.keep_with_next = True

        r_clbl = p_cmd.add_run("Reference Commands:  ")
        r_clbl.font.size = Pt(9)
        r_clbl.font.bold = True
        r_clbl.font.color.rgb = hex_to_rgb(COLOR_TEXT_MUTED)

        for cmd in commands:
            r_c = p_cmd.add_run(f" {cmd} ")
            r_c.font.name = "Consolas"
            r_c.font.bold = True
            r_c.font.size = Pt(9)
            r_c.font.color.rgb = hex_to_rgb(COLOR_PRIMARY)
            p_cmd.add_run("  ")

    def add_question_box(q_num: int, q_text: str) -> None:
        # Question container table: 3 rows (Header, Command, Screenshot)
        q_tbl = doc.add_table(rows=3, cols=1)
        q_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        q_tbl.autofit = False
        set_table_padding(q_tbl, top=100, bottom=100, left=140, right=140)

        # Borders: subtle card border
        tblPr = q_tbl._tbl.tblPr
        tblBorders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>\n'
            f'  <w:top w:val="single" w:sz="6" w:space="0" w:color="{COLOR_BORDER_LIGHT}"/>\n'
            f'  <w:bottom w:val="single" w:sz="10" w:space="0" w:color="{COLOR_BORDER}"/>\n'
            f'  <w:left w:val="single" w:sz="6" w:space="0" w:color="{COLOR_BORDER_LIGHT}"/>\n'
            f'  <w:right w:val="single" w:sz="6" w:space="0" w:color="{COLOR_BORDER_LIGHT}"/>\n'
            f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{COLOR_BORDER_LIGHT}"/>\n'
            f'  <w:insideV w:val="none"/>\n'
            f'</w:tblBorders>'
        )
        tblPr.append(tblBorders)

        # Row 0: Question Header
        r0 = q_tbl.rows[0]
        make_row_cant_split(r0)
        c0 = r0.cells[0]
        c0.width = Inches(6.8)
        set_cell_shading(c0, COLOR_SURFACE)
        set_cell_padding(c0, top=120, bottom=120, left=140, right=140)

        p0 = c0.paragraphs[0]
        p0.paragraph_format.space_before = Pt(0)
        p0.paragraph_format.space_after = Pt(0)
        p0.paragraph_format.keep_with_next = True

        r_qn = p0.add_run(f"Question {q_num}: ")
        r_qn.font.bold = True
        r_qn.font.size = Pt(10)
        r_qn.font.color.rgb = hex_to_rgb(COLOR_ACCENT)

        r_qt = p0.add_run(q_text)
        r_qt.font.bold = True
        r_qt.font.size = Pt(10)
        r_qt.font.color.rgb = hex_to_rgb(COLOR_PRIMARY)

        # Row 1: Command Area
        r1 = q_tbl.rows[1]
        make_row_cant_split(r1)
        c1 = r1.cells[0]
        c1.width = Inches(6.8)
        set_cell_shading(c1, "0D1117")
        set_cell_padding(c1, top=90, bottom=90, left=140, right=140)

        p1 = c1.paragraphs[0]
        p1.paragraph_format.space_before = Pt(0)
        p1.paragraph_format.space_after = Pt(0)
        p1.paragraph_format.keep_with_next = True

        r_prmpt = p1.add_run("$ ")
        r_prmpt.font.name = "Consolas"
        r_prmpt.font.bold = True
        r_prmpt.font.size = Pt(10.5)
        r_prmpt.font.color.rgb = hex_to_rgb("3FB950")

        # Row 2: Screenshot Drop Zone (Dashed Border + Camera Glyph)
        r2 = q_tbl.rows[2]
        make_row_cant_split(r2)
        c2 = r2.cells[0]
        c2.width = Inches(6.8)
        set_cell_shading(c2, "FAFBFD")
        set_cell_padding(c2, top=130, bottom=130, left=140, right=140)

        # Set dashed border on screenshot container
        tcPr2 = c2._tc.get_or_add_tcPr()
        borders2 = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>\n'
            f'  <w:top w:val="dashed" w:sz="6" w:space="0" w:color="9AA4B2"/>\n'
            f'  <w:left w:val="dashed" w:sz="6" w:space="0" w:color="9AA4B2"/>\n'
            f'  <w:bottom w:val="dashed" w:sz="6" w:space="0" w:color="9AA4B2"/>\n'
            f'  <w:right w:val="dashed" w:sz="6" w:space="0" w:color="9AA4B2"/>\n'
            f'</w:tcBorders>'
        )
        tcPr2.append(borders2)

        # Set minimum height for screenshot box (1600 twips = 1.11 inches, atLeast)
        trPr2 = r2._tr.get_or_add_trPr()
        trPr2.append(parse_xml(f'<w:trHeight {nsdecls("w")} w:val="1600" w:hRule="atLeast"/>'))

        p2 = c2.paragraphs[0]
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(2)
        r_slbl = p2.add_run("EXECUTION SCREENSHOT:")
        r_slbl.font.size = Pt(8)
        r_slbl.font.bold = True
        r_slbl.font.color.rgb = hex_to_rgb(COLOR_TEXT_MUTED)

        p2_hint = c2.add_paragraph()
        p2_hint.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2_hint.paragraph_format.space_before = Pt(14)
        p2_hint.paragraph_format.space_after = Pt(14)
        r_hint = p2_hint.add_run("📷  Paste terminal screenshot here (must show executed command and output)")
        r_hint.font.italic = True
        r_hint.font.size = Pt(8.5)
        r_hint.font.color.rgb = hex_to_rgb("8B949E")

        # Spacing after question table
        p_post = doc.add_paragraph()
        p_post.paragraph_format.space_before = Pt(0)
        p_post.paragraph_format.space_after = Pt(4)

    # -----------------------------------------------------------------------
    # POPULATING ALL 13 QUESTIONS (5-Page Structured Layout)
    # -----------------------------------------------------------------------

    # PAGE 2+: Natural semantic flow with cantSplit and keep_with_next
    # Section 1
    add_section_header("SECTION 01  •  Basic File and Directory Commands", ["ls", "cd", "pwd", "mkdir", "rmdir", "rm"])
    add_question_box(1, "How would you list all files, including hidden ones, in a directory?")
    add_question_box(2, "What command would you use to navigate to the parent directory?")
    add_question_box(3, "How can you create a new directory named projects?")

    # Section 2
    add_section_header("SECTION 02  •  File Manipulation Commands", ["cp", "mv", "cat", "nano", "touch"])
    add_question_box(4, "How would you copy a file named report.txt to the /backup directory?")
    add_question_box(5, "What command would you use to rename draft.txt to final.txt?")
    add_question_box(6, "How can you create a new empty file named log.txt?")

    # Section 3
    add_section_header("SECTION 03  •  Viewing and Searching Files", ["less", "grep", "find"])
    add_question_box(7, 'How would you search for the word "success" in a file named results.txt?')
    add_question_box(8, "What command would you use to view the contents of manual.pdf one page at a time?")
    add_question_box(9, "How can you find all .jpg files in the /pictures directory?")

    # Section 4 & Section 5
    add_section_header("SECTION 04  •  Process Management: Viewing Processes", ["ps", "top"])
    add_question_box(10, "How would you list all currently running processes?")
    add_question_box(11, "What command provides a dynamic real-time view of running processes?")

    add_section_header("SECTION 05  •  Process Management: Managing Processes", ["kill", "killall", "pkill"])
    add_question_box(12, "How would you terminate a process with PID 5678?")
    add_question_box(13, "What command would you use to kill all instances of chrome?")

    # Final Completion Callout
    end_tbl = doc.add_table(rows=1, cols=1)
    end_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    end_tbl.autofit = False
    c_end = end_tbl.cell(0, 0)
    c_end.width = Inches(6.8)
    set_cell_shading(c_end, COLOR_SURFACE)
    set_cell_padding(c_end, top=140, bottom=140, left=180, right=160)

    tcPr_end = c_end._tc.get_or_add_tcPr()
    borders_end = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>\n'
        f'  <w:top w:val="single" w:sz="12" w:space="0" w:color="{COLOR_ACCENT}"/>\n'
        f'  <w:left w:val="single" w:sz="6" w:space="0" w:color="{COLOR_BORDER_LIGHT}"/>\n'
        f'  <w:bottom w:val="single" w:sz="6" w:space="0" w:color="{COLOR_BORDER_LIGHT}"/>\n'
        f'  <w:right w:val="single" w:sz="6" w:space="0" w:color="{COLOR_BORDER_LIGHT}"/>\n'
        f'</w:tcBorders>'
    )
    tcPr_end.append(borders_end)
    make_row_cant_split(end_tbl.rows[0])

    p_end = c_end.paragraphs[0]
    p_end.paragraph_format.space_before = Pt(0)
    p_end.paragraph_format.space_after = Pt(2)
    r_etitle = p_end.add_run("SUBMISSION VERIFICATION CHECKLIST\n")
    r_etitle.font.bold = True
    r_etitle.font.size = Pt(9)
    r_etitle.font.color.rgb = hex_to_rgb(COLOR_ACCENT)

    r_ebody = p_end.add_run(
        "✓ All 13 questions have exact command syntax typed.\n"
        "✓ Screenshots clearly show the terminal prompt, executed command, and output.\n"
        "✓ Student names and university IDs on the cover page are verified before final submission."
    )
    r_ebody.font.size = Pt(9)
    r_ebody.font.color.rgb = hex_to_rgb(COLOR_TEXT)

    # Save DOCX
    doc.save(output_path)
    print(f"Successfully saved redesigned document to {output_path} ({os.path.getsize(output_path)} bytes)")
    return doc

if __name__ == "__main__":
    out_docx = "/home/albraa/Desktop/CS345_Linux_Commands_Lab_Template.docx"
    logo = "/home/albraa/Desktop/CS345/assets/taibah_logo.png"
    build_redesigned_document(out_docx, logo)
