"""
test_docx_craft.py - Verification & Showcase Test for docx_craft.py

Generates a realistic, multi-page professional document using Taibah University
Cybersecurity curriculum data (from data/courses.json), renders it to PDF via
headless LibreOffice, and generates high-resolution page PNGs via pdftoppm.
"""

import json
import os
import sys
from pathlib import Path
from docx import Document

# Add scripts directory to sys.path
SCRIPT_DIR = Path(__file__).resolve().parent
sys.path.insert(0, str(SCRIPT_DIR))

from docx_craft import (
    PALETTES,
    apply_page_geometry,
    setup_document_styles,
    add_cover_page,
    setup_header_footer,
    add_heading,
    create_styled_table,
    add_callout,
    add_stat_card_row,
    render_and_preview,
)

def run_test():
    repo_root = SCRIPT_DIR.parents[3]
    courses_json_path = repo_root / "data" / "courses.json"
    assert courses_json_path.exists(), f"Could not find courses.json at {courses_json_path}"

    with open(courses_json_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    program = data.get("program", {})
    courses = data.get("courses", [])

    print(f"Loaded program: {program.get('name')} from {program.get('university')}")
    print(f"Total courses loaded: {len(courses)}")

    doc = Document()
    palette_name = "slate_executive"

    # 1. Page Geometry (Executive 0.75" / 0.8" margins)
    apply_page_geometry(doc.sections[0], preset="executive")

    # 2. Typography Setup
    palette = setup_document_styles(doc, palette_name=palette_name, font_preset="modern")

    # 3. Cover Page
    add_cover_page(
        doc,
        title=f"{program.get('name', 'Cybersecurity')} Curriculum Plan",
        subtitle=f"Degree Roadmap, Course Distribution & Academic Milestones ({program.get('durationYears', 5)}-Year Program)",
        author_org=f"{program.get('university', 'Taibah University')} · Department of Computer Science",
        date_version="Academic Year 2026–2027 · Curriculum Specification v2.1",
        archetype="executive_stripe",
        palette_name=palette_name,
    )

    # 4. Header / Footer with dynamic page numbers
    setup_header_footer(
        doc,
        title_left="Taibah University · Cybersecurity Curriculum",
        doc_info_right="Official Degree Plan",
        palette_name=palette_name,
        show_page_numbers=True,
        suppress_cover_page=True,
    )

    # 5. Section 1: Executive Overview
    add_heading(doc, "1. Executive Overview & Program Metrics", level=1, subtitle="Key figures and degree architecture at a glance")

    total_credits = sum(c.get("credits", 0) for c in courses)
    add_stat_card_row(doc, [
        (str(program.get("terms", 10)), "Academic Terms", "5 full academic years"),
        (str(total_credits), "Total Credits", "Accredited degree"),
        (str(len(courses)), "Total Courses", "Core + Elective"),
        ("Level 7", "NQF Classification", "Bachelor of Science"),
    ], palette_name=palette_name)

    p_intro = doc.add_paragraph()
    p_intro.add_run(
        "This document specifies the complete course sequencing and prerequisite structure for the Bachelor of Science "
        "in Cybersecurity at Taibah University. Students must fulfill all foundational science and mathematics requirements "
        "before entering advanced offensive security, cryptography, and digital forensics modules."
    )

    # 6. Callout Alert
    add_callout(
        doc,
        text="All foundational mathematics (MATH 101) and programming prerequisites must be completed with a grade of 'C' or higher before enrolling in third-year network security and cryptography laboratories.",
        title="Prerequisite Policy Advisory",
        kind="warning",
        palette_name=palette_name,
    )

    # 7. Section 2: Year 1 & 2 Foundational Courses Table
    add_heading(doc, "2. Preparatory & Foundational Course Distribution", level=1, subtitle="First-year and second-year foundational sequence")

    term_1_and_2 = [c for c in courses if c.get("term", 0) in (1, 2)]
    headers = ["Course Code", "Course Title", "Term", "Credits", "Category"]
    rows = []
    for c in term_1_and_2:
        rows.append([
            c.get("code", ""),
            c.get("name", ""),
            f"Term {c.get('term', '')}",
            f"{c.get('credits', 0)} cr",
            c.get("type", "Plan"),
        ])

    create_styled_table(
        doc,
        headers=headers,
        rows=rows,
        col_widths=[1.2, 3.2, 0.9, 0.8, 0.9],
        palette_name=palette_name,
        zebra_striping=True,
    )

    # 8. Section 3: Advanced Specialization Overview
    add_heading(doc, "3. Core Security Modules & Capstone Distribution", level=1, subtitle="Terms 3 through 10 milestone breakdown")

    add_callout(
        doc,
        text="The final academic year focuses on Senior Design Project I & II, alongside an intensive Cooperative Training internship in an enterprise security operations center (SOC) or defensive security engineering role.",
        title="Graduation Milestone Note",
        kind="takeaway",
        palette_name=palette_name,
    )

    # Later courses sample
    advanced_courses = [c for c in courses if c.get("term", 0) >= 3][:8]
    adv_rows = []
    for c in advanced_courses:
        adv_rows.append([
            c.get("code", ""),
            c.get("name", ""),
            f"Term {c.get('term', '')}",
            f"{c.get('credits', 0)} cr",
            c.get("type", "Plan"),
        ])

    create_styled_table(
        doc,
        headers=headers,
        rows=adv_rows,
        col_widths=[1.2, 3.2, 0.9, 0.8, 0.9],
        palette_name=palette_name,
        zebra_striping=True,
    )

    # Save output DOCX
    test_out_dir = SCRIPT_DIR / "test_output"
    test_out_dir.mkdir(parents=True, exist_ok=True)
    docx_path = test_out_dir / "cybersecurity_plan_showcase.docx"
    doc.save(docx_path)
    print(f"Successfully created DOCX: {docx_path} ({docx_path.stat().st_size} bytes)")
    assert docx_path.exists() and docx_path.stat().st_size > 5000, "DOCX generation failed"

    # Test PDF and PNG rendering
    print("Converting DOCX to PDF via LibreOffice headless...")
    pdf_path, images = render_and_preview(docx_path, output_dir=test_out_dir)

    print(f"Successfully generated PDF: {pdf_path} ({pdf_path.stat().st_size} bytes)")
    assert pdf_path.exists() and pdf_path.stat().st_size > 5000, "PDF rendering failed"

    print(f"Successfully generated {len(images)} preview images:")
    for img in images:
        print(f"  - {img.name} ({img.stat().st_size} bytes)")
    assert len(images) >= 2, f"Expected at least 2 pages, got {len(images)}"

    print("\nALL ASSERTIONS PASSED! DOCX design craft test completed successfully.")
    return docx_path, pdf_path, images

if __name__ == "__main__":
    run_test()
